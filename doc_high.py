from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import sys
import cloudconvert


api = cloudconvert.Api(
    'tggDTnhhQ2ogvUSfN0TpZASfNS5Pq5ZbqfE8EuQVLQeRjAq0uVNBX4keX3O1FyeR')
process = api.convert({
    'inputformat': 'txt',
    'outputformat': 'docx',
    'input': 'upload',
    'file': open("test3.txt", 'rb')
})
process.wait()  # wait until conversion finished
pdf_file_name = str("test3")+".docx"
process.download(pdf_file_name)
doc = Document("test3.docx")

for para in doc.paragraphs:
    start = para.text.find("bh")
    print(start)
    if start > -1:
        pre = para.text[:start]
        print(pre)
        post = para.text[start+len("bh"):]
        print(post)
        para.text = pre
        para.add_run("bh")
        para.runs[1].font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        para.add_run(post)

start = 10
if start > -1:
    pre = "hvg bh hkjbhjj"
    post = "bhbhbhtest3"
    para.text = pre
    para.add_run("bh")
    para.runs[1].font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
    para.add_run(post)


doc.save("fuck.docx")
